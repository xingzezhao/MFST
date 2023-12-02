#!/usr/bin/env python
# coding: utf-8

import configparser
import MDAnalysis as mda
import numpy as np
import os
import matplotlib.pyplot as plt
import subprocess
from docx import Document
import scipy
from scipy import optimize as opt
from datetime import datetime
import glob

def read_config(config_file):
    # 创建配置解析器对象
    config = configparser.ConfigParser()
    MFST_path = os.path.abspath('.')
#     print(MFST_path)
    # 读取配置文件
    config.read(config_file)
    for section in config.sections():
        for option, value in config.items(section):
            if value[0] == '.':
                config.set(section, option, os.path.abspath(value))
    # 获取配置项的值
    config_values = {
        'no_surfts': config.getint('pyModeling input', 'no_surfts'),
        'surfts_type': config.get('pyModeling input', 'surfts_type').split(),
        'interfc_type': config.get('pyModeling input', 'interfc_type'),
        'itf_surfts': [int(value) for value in config.get('pyModeling input', 'itf_surfts').split()],
        'waterbox_size_list': [float(value) for value in config.get('pyModeling input', 'waterbox_size').split()],
        'simulationbox_list': [float(value) for value in config.get('pyModeling input', 'simulationbox_size').split()],
        'simulationsys_name': config.get('pyModeling input', 'simulationsys_name'),
        'temp': config.getfloat('pySimulating input', 'Temp'),
        'packmolpath': config.get('pyModeling input', 'packmolpath'),
        'WFF': config.get('pySimulating input', 'SimuWatMod'),
        'eqnsteps':  config.getint('pySimulating input', 'eqnsteps'),
        'simnsteps':  config.getint('pySimulating input', 'simnsteps'),
        'simdt': config.getfloat('pySimulating input', 'dt'),
        'outflames': config.getint('pySimulating input', 'outflames'),
        'rcoulomb': config.getfloat('pySimulating input', 'rcoulomb'),
        'rlist': config.getfloat('pySimulating input', 'rlist'),
        'rvdw': config.getfloat('pySimulating input', 'rvdw'),
        'gmxpath': config.get('pySimulating input', 'gmxpath'),
        'pions': config.get('pySimulating input', 'pions'),
        'nions': config.get('pySimulating input', 'nions'),
        'conc': config.getfloat('pySimulating input', 'conc'),
        'emtol': config.get('pySimulating input', 'emtol'),
        'jcrls': config.get('pySimulating input', 'jcrls'),
    }
    
    jcrtls = config_values['jcrls']
    if jcrtls == 'Slurm':
        config_values['jcspath'] = config.get('pySimulating input', 'sltmpath')
        config_values['jcline'] = 'sbatch'
    elif jcrtls == 'PBS':
        config_values['jcspath'] = config.get('pySimulating input', 'pbstmpath')
        config_values['jcline'] = 'qsub'
    else:
        print(f'ERROR: Cannot recognize the Job Crtl Sys:{jcrtls}, you should use Slurm or PBS')

    return config_values

def find_workdir(pymdlog,simulationsys_name):
    with open(pymdlog, 'r') as log_file:
        # 遍历文件的每一行
        for line in log_file:
            if 'Project Name' in line:
                if line.split(': ')[-1].strip() == simulationsys_name :
                    # 检查是否包含目录信息
                    working_directory_line = next(log_file).strip()
                    if 'Working Directory' in working_directory_line:
                        # 提取目录路径
                        directory_path = working_directory_line.split(': ')[-1].strip()

                        # 切换工作目录
                        os.chdir(directory_path)
                        print(f'Changed working directory to: {directory_path}')
                        break  # 读到目录信息后退出循环
                
def energyanl(simpart, config_values):
    sn = config_values['simulationsys_name']
    engfl = sn+f'_{simpart}.edr'
    engout = sn+f'_{simpart}_eng.xvg'
    gmx = config_values['gmxpath']

    engln = [gmx, 'energy', '-f', engfl, '-o', engout]
    process = subprocess.Popen(engln, stdin=subprocess.PIPE, text=True)
    process.communicate(input='Potential\n')
    process.wait()

    engoutdata = np.loadtxt(engout,comments=['@','#'])
    engoutdata = np.transpose(engoutdata)
    plt.plot(engoutdata[0],engoutdata[1], color='#f01138', linewidth=2.5)
    plt.xlabel('time (ps)')
    plt.ylabel('energy (kJ/mol)')
    plt.ticklabel_format(style='sci', axis='y', scilimits=(0,0))
    plt.savefig(sn+f'_{simpart}.jpg', dpi=600)
    plt.close()
    deng = engoutdata[1][round(len(engoutdata[1])/2):].max() - engoutdata[1][round(len(engoutdata[1])/2):].min()
    return deng

def split_HTGroup(config_values):
    sn = config_values['simulationsys_name']
    u = mda.Universe(sn+'_em.gro',sn+'_em.gro')
    surf = []
    counter = 0
    for i in config_values['surfts_type']:
        snum = config_values['itf_surfts'][counter]*2
        surftot = u.select_atoms(f'resname {i}')
        surf.append(surftot[:int(len(surftot)/snum)])
        counter += 1

    coTl = []
    coHl = []
    Tgl = []
    Hgl = []
    for surf in surf: 
        carbon_atoms = surf.select_atoms('name C* and not name H*')
        other_atoms = surf.select_atoms('not name C* and not name H*')
        # 初始化变量
        max_distance = 0
        farthest_atom_pair = None

        # 遍历所有碳原子对
        for carbon_atom in carbon_atoms:
                for other_atom in other_atoms:
                    # 计算两个原子之间的距离
                    distance = np.linalg.norm(carbon_atom.position - other_atom.position)

                    # 更新最大距离和原子对
                    if distance > max_distance:
                        max_distance = distance
                        farthest_atom_pair = (carbon_atom, other_atom)

        if farthest_atom_pair is not None:
            carbon_atom_index = farthest_atom_pair[0].index
            other_atom_index = farthest_atom_pair[1].index
        else:
            print("Error: No hydrophobic tail atoms found.")


        coT = farthest_atom_pair[0]
        coH = farthest_atom_pair[1]
        Tgroup  = []
        for atom in surf:
            distance = np.linalg.norm(atom.position - coT.position)
            if distance < max_distance/2:
                Tgroup.append(atom.name)
        Hgroup  = []
        for atom in surf:
            distance = np.linalg.norm(atom.position - coH.position)
            if distance < max_distance/2:
                Hgroup.append(atom.name)
        coT = coT.name
        coH = coH.name
        coTl.append(coT)
        coHl.append(coH)
        Tgl.append(Tgroup)
        Hgl.append(Hgroup)
    return coTl,coHl,Tgl, Hgl


def getZdensity(positions,config_values,M):
    Zmax = config_values['simulationbox_list'][2]* 10
    Zsl = 400
    dz = Zmax / Zsl
    xlist = np.arange(0,Zmax+dz,dz)
    positions = np.transpose(positions)[-1]
    counter = np.histogram(positions, xlist)[0]
    rho = counter*(M*400/(0.6022*config_values['simulationbox_list'][0]*config_values['simulationbox_list'][1]*config_values['simulationbox_list'][2]))
    
    return rho

def sigmoid(X, a, b, c):
    return c / (1 + a * np.exp(b * X))

def curve_fit(testlist, c):
    A, B, c = opt.curve_fit(sigmoid, testlist[0], testlist[1], maxfev = 10000)[0]
    fitx = np.arange(min(testlist[0]),max(testlist[0]),0.0001)
    fity = sigmoid(fitx,A, B, c )
    return fitx, fity

def pyAnalysing_main():
    config_file_path = './MFST.rc'  
    config_values = read_config(config_file_path)
    MFST_path = os.path.abspath('.')
    find_workdir('./pyModeling.log',config_values['simulationsys_name'])
    sn = config_values['simulationsys_name']


    ## simulation info
    ## simulation name sn; simulation type st; surfts name; surfts number sufn; simulation box size sbs; thickness of water tow;
    ## no of water wnum; ions; simulation steps ss; temp
    surfts = ''
    for i in config_values['surfts_type']:
        surfts += i

    if config_values['interfc_type'] == 'GLG':
        st = 'Gas Liquid Gas'
    elif config_values['interfc_type'] == 'OLO':
        st = 'Oil Liquid Oil'
    elif config_values['interfc_type'] == 'L':
        st = 'pure Liquid'

    sufn = str(config_values['itf_surfts'][0]) + ' in each sides'

    sbs = str(config_values['simulationbox_list'][0]) + ' x ' + str(config_values['simulationbox_list'][1]) + ' x ' + str(config_values['simulationbox_list'][2]) + ' nm'
    tow = str(config_values['waterbox_size_list'][2]) + ' nm'
    with open(sn+'.top') as file:
        lines = file.readlines()
    for line in lines:
        if 'SOL' in line:
            wnum = line.split()[-1]

    ions = str(config_values['conc']) +' M '+config_values['pions'] + config_values['nions']

    ss = str(config_values['eqnsteps'] * config_values['simdt']) + ' ps '+'equilibrium & '+ str(config_values['simnsteps'] * config_values['simdt']) + ' ps '+'MD'
    temp = str(config_values['temp']) + ' K'

    SIout = """This simulation system is about:
    simulation name : {sn}
    simulation type : {st}
    surfactants info : {surfts} {sufn}
    simulation box size : {sbs}
    thickness of water : {tow}
    number of water : {wnum}
    ions : {ions}
    simulation steps : {ss}
    temperature : {temp}
    """
    SIout1 = SIout.format(
        sn = sn,
        st = st,
        surfts = surfts,
        sufn = sufn,
        sbs = sbs,
        tow = tow,
        wnum = wnum,
        ions = ions,
        ss = ss,
        temp = temp
    )

    ## energy 
    eqdeng = energyanl('eq',config_values)
    mddeng = energyanl('md',config_values)
    ENout = f"""the average of energy in equilibrium is {eqdeng} kJ/mol
    the average of energy in equilibrium is {mddeng} kJ/mol"""

    ## interface

    ### DENSITY
    coTl,coHl,Tgl, Hgl = split_HTGroup(config_values)
    noST = len(coTl)
    u = mda.Universe(sn+'_md.gro',sn+'_md.trr')

    u2 = mda.Universe('./water.pdb')
    u2 = u2.select_atoms('name *')
    waterms = sum(u2.masses)
    for i in range(noST):
        surfn = config_values['surfts_type'][i]
        u99 = mda.Universe(config_values['simulationsys_name']+'_em.gro')
        surftot = u99.select_atoms(f'resname {surfn}')
        u1 = u99.select_atoms(f'resname {surfn}')[:int(len(surftot)/(config_values['itf_surfts'][i]*2))]
        u1Tot = u1.select_atoms('name *')
        u1T = u1.select_atoms(f"name {' '.join(Tgl[i])}")
        u1H = u1.select_atoms(f"name {' '.join(Hgl[i])}")
        surfmas = sum(u1Tot.masses)
        Tmas = sum(u1T.masses)
        Hmas = sum(u1H.masses)
        u1Tc = u1T.center_of_mass()
        u1Tp = u1T.positions
        minindex = np.argmin(np.linalg.norm(u1Tp-u1Tc,axis=1))
        u1Tc = u1T[minindex].name

        u1Hc = u1H.center_of_mass()
        u1Hp = u1H.positions
        minindex = np.argmin(np.linalg.norm(u1Hp-u1Hc,axis=1))
        u1Hc = u1H[minindex].name
        Sgclst = []
        Tgclst = []
        Hgclst = []
        for flame in u.trajectory:
            coT = u.select_atoms(f'resname {surfn} and name {coTl[i]}').positions
            coH = u.select_atoms(f'resname {surfn} and name {coHl[i]}').positions
            Tg = u.select_atoms(f"resname {surfn} and name {u1Tc}").positions
            Hg = u.select_atoms(f"resname {surfn} and name {u1Hc}").positions
            Sg = u.select_atoms(f"resname SOL and name O*").positions
            Sgc = getZdensity(Sg,config_values,waterms)
            Tgc = getZdensity(Tg,config_values,Tmas)
            Hgc = getZdensity(Hg,config_values,Hmas)
            Sgclst.append(Sgc)
            Tgclst.append(Tgc)
            Hgclst.append(Hgc)
        Sgclstm = np.mean(Sgclst,axis=0)
        Tgclstm = np.mean(Tgclst,axis=0)
        Hgclstm = np.mean(Hgclst,axis=0)
        Zmax = config_values['simulationbox_list'][2]
        Zsl = 400
        dz = Zmax / Zsl
        densityxlist = np.arange(0,Zmax,dz)
        plt.plot(densityxlist,Sgclstm,linewidth=2.5,color='b',label='Water')
        plt.plot(densityxlist,Tgclstm,linewidth=2.5,color='#800080',label=f'{surfn} Tail')    
        plt.plot(densityxlist,Hgclstm,linewidth=2.5,color='#FF0000',label=f'{surfn} Head')    
        plt.xlabel('Z coordinate (nm)')
        plt.ylabel('Density (g/L)')
        plt.legend(frameon=False)
        plt.savefig(sn+f'_{surfn}density.jpg',dpi=600)
        plt.close()

    ### GDS & Eqth
    y1 = Sgclstm[:round(len(Sgclstm)/2)]
    y2 = Sgclstm[round(len(Sgclstm)/2):]
    xlist = np.arange(0,Zmax,dz)
    x1 = xlist[:round(len(xlist)/2)]
    x2 = xlist[round(len(xlist)/2):]
    x1fit, y1fit = curve_fit([x1, y1], max(y1))
    x2fit, y2fit = curve_fit([x2-10, y2[::-1]], max(y2))
    x2fit = x2fit+10
    y2fit = y2fit[::-1]

    ndx901 = np.argmin(abs(y1fit-max(y1fit)*0.9))
    ndx101 = np.argmin(abs(y1fit-max(y1fit)*0.1))
    ndx902 = np.argmin(abs(y2fit-max(y2fit)*0.9))
    ndx102 = np.argmin(abs(y2fit-max(y2fit)*0.1))

    GDS = np.mean([x1fit[ndx901] - x1fit[ndx101],x2fit[ndx102] - x2fit[ndx902]])
    eqth = x2fit[ndx902] - x1fit[ndx901]

    GDSout = f'GDS aka. Gibbs dividing surface of This water interface is : %.5f nm'%GDS
    EQTHout = f'effective film thickness aka. 90-90 thickness of This water interface is : %.5f nm'%eqth

    ### interface tension GAMMA

    sn = config_values['simulationsys_name']
    engfl = sn+f'_md.edr'
    engout = sn+f'_md_ift.xvg'
    gmx = config_values['gmxpath']

    engln = [gmx, 'energy', '-f', engfl, '-o', engout]
    process = subprocess.Popen(engln, stdin=subprocess.PIPE, text=True)
    process.communicate(input='#Surf*SurfTen\n')
    process.wait()

    engoutdata = np.loadtxt(engout,comments=['@','#'])
    engoutdata = np.transpose(engoutdata)

    iftm = np.mean(engoutdata[-1]/20)

    IFTout = f'interface tension of this interface is : {iftm} mN/m'

    ## H-Bond
    sn = config_values['simulationsys_name']
    stufl = sn+f'_md.tpr'
    engfl = sn+f'_md.trr'
    gmx = config_values['gmxpath']
    HBout = ''
    for surf in config_values['surfts_type']:
        engout = sn+f'_md_hb_{surf}.xvg'
        engln = [gmx, 'hbond','-s', stufl, '-f', engfl, '-num', engout]
        process = subprocess.Popen(engln, stdin=subprocess.PIPE, text=True)
        process.communicate(input=f'{surf} SOL')
        process.wait()
        engoutdata = np.loadtxt(engout,comments=['@','#'])
        engoutdata = np.transpose(engoutdata)
        hbnumm = np.mean(engoutdata[1])
        HBout += f'the HBonds number between SOL and {surf} is : {hbnumm}'
        hbxlist = engoutdata[0]/1000
        plt.plot(hbxlist, engoutdata[1], linewidth=2.5,label=f'{surf}')

    plt.xlabel('time (ns)')
    plt.ylabel('No. of HBonds')
    plt.savefig(sn+f'_HB.jpg',dpi=600)
    plt.close()

    ## MSD of Water

    sn = config_values['simulationsys_name']
    stufl = sn+f'_md.tpr'
    engfl = sn+f'_md.trr'
    gmx = config_values['gmxpath']
    HBout = ''

    engout = sn+f'_msd.xvg'
    engln = [gmx, 'msd','-s', stufl, '-f', engfl, '-o', engout]
    process = subprocess.Popen(engln, stdin=subprocess.PIPE, text=True)
    process.communicate(input=f'SOL')
    process.wait()
    engoutdata = np.loadtxt(engout,comments=['@','#'])
    engoutdata = np.transpose(engoutdata)
    msdxlist = engoutdata[0]/1000
    plt.plot(msdxlist, engoutdata[1], linewidth=2.5)

    plt.xlabel('time (ns)')
    plt.ylabel('MSD')
    plt.savefig(sn+f'_MSD.jpg',dpi=600)
    plt.close()


    ## ORDER PARAMETER
    coTl,coHl,Tgl, Hgl = split_HTGroup(config_values)
    noST = len(coTl)
    u = mda.Universe(sn+'_md.gro',sn+'_md.trr')

    u2 = mda.Universe('./water.pdb')
    u2 = u2.select_atoms('name *')
    waterms = sum(u2.masses)
    opOUT = ''
    header_comment = "### The Order Parameter, every line is surf in one flame"
    for i in range(noST):
        surfn = config_values['surfts_type'][i]
        flamelst = []
        for flame in u.trajectory:
            coT = u.select_atoms(f'resname {surfn} and name {coTl[i]}').positions
            coH = u.select_atoms(f'resname {surfn} and name {coHl[i]}').positions
            opbox = []
            for f in range(0,len(coT)):
                newloc = coT[f]-coH[f]
                op = (3* ((abs(newloc[2] / np.sqrt(newloc[0]**2 + newloc[1]**2 + newloc[2]**2)))**2) -1)/2
                opbox.append(op)
            flamelst.append(opbox)
        surfop = np.mean(flamelst)
        opOUT += f'The mean Order Parameter of {surfn} is : {surfop}'
        flamelst = np.array(flamelst)
        np.savetxt(sn + f'{surfn}_op.xvg',flamelst,header=header_comment)


    ## WRITE THE DOCX
    doc = Document()

    doc.add_heading('Molecular Dynamic Simulation Report', level=1)
    current_datetime = datetime.now()
    doc.add_paragraph(f'{current_datetime}')
    doc.add_paragraph(f'\n')
    doc.add_heading('Part 1. Simulation Info', level=2)
    doc.add_paragraph(SIout1)
    doc.add_heading('Part 2. Energy', level=2)
    doc.add_paragraph(ENout)
    doc.add_picture(f'{sn}_eq.jpg')
    doc.add_paragraph('fig1 the energy of system during equilibrium')
    doc.add_picture(f'{sn}_md.jpg')
    doc.add_paragraph('fig2 the energy of system during MD')
    doc.add_paragraph(f'\n')
    doc.add_heading('Part 3. Interface', level=2)
    for surf in config_values['surfts_type']:
        doc.add_picture(f'./{sn}_{surf}density.jpg')
    doc.add_paragraph('fig3 the density profile')
    doc.add_paragraph(GDSout)
    doc.add_paragraph(EQTHout)
    doc.add_paragraph(IFTout)
    doc.add_heading('Part 4. H-bonds', level=2)
    doc.add_picture(f'{sn}_HB.jpg')
    doc.add_paragraph('fig4 the number of H-bonds')
    doc.add_paragraph(HBout)
    doc.add_heading('Part 5. MSD', level=2)
    doc.add_picture(f'{sn}_MSD.jpg')
    doc.add_paragraph('fig5 the MSD of Water')
    doc.add_heading('Part 6. Order Parameter', level=2)
    # doc.add_picture(f'{sn}_op.jpg')
    # doc.add_paragraph('fig6 the Order Parameter of Surf')
    doc.add_paragraph(opOUT)

    doc.add_paragraph('\n\n\n\n')
    doc.add_paragraph('MFST MolFilmStabTool, is a a full-process automated foam film stability research software')
    doc.add_paragraph('v 1.0')
    doc.add_paragraph('code by Xingze Zhao, you can connect with me by email: xingzezhao@gmail.com or by Github: https://github.com/xingzezhao')
    doc.add_paragraph('MFST © 2023 MTSD@UPC')
    doc.save(f'{sn}_report.docx')
    os.chdir(MFST_path)
    
if __name__ == "__main__":
    pyAnalysing_main()